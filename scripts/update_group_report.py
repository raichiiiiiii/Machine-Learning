from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from sklearn.metrics import precision_recall_curve, roc_curve


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
FIG_DIR = REPORTS / "group_report_figures"
SOURCE = REPORTS / "ML_Group_J_Full_Report_source.docx"
OUTPUT = REPORTS / "ML_Group_J_Full_Report_UPDATED.docx"
FIG_DIR.mkdir(parents=True, exist_ok=True)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def replace_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def insert_image_at_placeholder(paragraph, image_path, width=5.8):
    image_paragraph = paragraph.insert_paragraph_before()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    remove_paragraph(paragraph)


def confusion_figure(model_name, tn, fp, fn, tp, out_path):
    cm = np.array([[tn, fp], [fn, tp]])
    fig, ax = plt.subplots(figsize=(4.6, 4.0))
    ax.imshow(cm, cmap="Blues")
    ax.set_title(f"{model_name} Confusion Matrix")
    ax.set_xlabel("Predicted label")
    ax.set_ylabel("Actual label")
    ax.set_xticks([0, 1], ["Non-fraud", "Fraud"])
    ax.set_yticks([0, 1], ["Non-fraud", "Fraud"])
    for i in range(2):
        for j in range(2):
            ax.text(j, i, f"{cm[i, j]:,}", ha="center", va="center", fontsize=11)
    fig.tight_layout()
    fig.savefig(out_path, dpi=180)
    plt.close(fig)


def generate_figures():
    results = pd.read_csv(ROOT / "model" / "Classical" / "csv" / "classical_results_comparison.csv")
    pca_results = results[results["dataset"] == "fraud_pca_95_variance"]

    for model, filename in [
        ("Random Forest", "figure_5_random_forest_confusion.png"),
        ("XGBoost", "figure_6_xgboost_confusion.png"),
    ]:
        row = pca_results[pca_results["model"] == model].iloc[0]
        confusion_figure(model, int(row.tn), int(row.fp), int(row.fn), int(row.tp), FIG_DIR / filename)

    all_rows = []
    mlp = pd.read_csv(ROOT / "model" / "MLP" / "csv" / "mlp_results_comparison.csv")
    rnn = pd.read_csv(ROOT / "model" / "RNN" / "csv" / "rnn_results_comparison.csv")
    all_rows.append({"model": "MLP", **mlp[mlp["dataset"] == "../PCA/fraud_pca_95_variance.csv"].iloc[0].to_dict()})
    all_rows.append({"model": "RNN", **rnn[rnn["dataset"] == "../PCA/fraud_pca_95_variance.csv"].iloc[0].to_dict()})
    for _, row in pca_results.iterrows():
        all_rows.append({"model": row["model"], "f1": row["f1"], "pr_auc": row["pr_auc"]})
    comparison = pd.DataFrame(all_rows).sort_values("pr_auc", ascending=True)

    colors = {
        "Random Forest": "#1F77B4",
        "XGBoost": "#2CA02C",
        "Decision Tree": "#FF7F0E",
        "MLP": "#8C564B",
        "RNN": "#9467BD",
        "Linear Regression": "#7F7F7F",
    }
    fig, axes = plt.subplots(1, 2, figsize=(10.8, 3.8))
    axes[0].barh(comparison["model"], comparison["pr_auc"], color=[colors[m] for m in comparison["model"]])
    axes[0].set_title("PR-AUC on PCA Feature Set")
    axes[0].set_xlabel("PR-AUC")
    axes[0].grid(axis="x", alpha=0.25)
    axes[1].barh(comparison["model"], comparison["f1"], color=[colors[m] for m in comparison["model"]])
    axes[1].set_title("F1-score on PCA Feature Set")
    axes[1].set_xlabel("F1-score")
    axes[1].grid(axis="x", alpha=0.25)
    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure_11_metric_comparison.png", dpi=180)
    plt.close(fig)

    mlp_pred = pd.read_csv(ROOT / "model" / "MLP" / "csv" / "fraud_pca_95_variance_mlp_test_predictions.csv")
    rnn_pred = pd.read_csv(ROOT / "model" / "RNN" / "csv" / "fraud_pca_95_variance_rnn_test_predictions.csv")
    classical_pred = pd.read_csv(ROOT / "model" / "Classical" / "csv" / "fraud_pca_95_variance_classical_test_predictions.csv")
    y_true = mlp_pred["y_true"].to_numpy()
    probs = {
        "MLP": mlp_pred["y_prob"].to_numpy(),
        "RNN": rnn_pred["y_prob"].to_numpy(),
        "Linear Regression": classical_pred["linear_regression_probability"].to_numpy(),
        "Decision Tree": classical_pred["decision_tree_probability"].to_numpy(),
        "Random Forest": classical_pred["random_forest_probability"].to_numpy(),
        "XGBoost": classical_pred["xgboost_probability"].to_numpy(),
    }
    fig, axes = plt.subplots(1, 2, figsize=(11, 3.8))
    for model, prob in probs.items():
        precision, recall, _ = precision_recall_curve(y_true, prob)
        fpr, tpr, _ = roc_curve(y_true, prob)
        axes[0].plot(recall, precision, label=model, linewidth=1.6)
        axes[1].plot(fpr, tpr, label=model, linewidth=1.6)
    axes[0].set_title("Precision-Recall Curves")
    axes[0].set_xlabel("Recall")
    axes[0].set_ylabel("Precision")
    axes[0].grid(alpha=0.25)
    axes[0].legend(fontsize=7)
    axes[1].plot([0, 1], [0, 1], "--", color="gray", linewidth=1)
    axes[1].set_title("ROC Curves")
    axes[1].set_xlabel("False Positive Rate")
    axes[1].set_ylabel("True Positive Rate")
    axes[1].grid(alpha=0.25)
    axes[1].legend(fontsize=7)
    fig.tight_layout()
    fig.savefig(FIG_DIR / "figure_12_roc_pr_curves.png", dpi=180)
    plt.close(fig)


def update_document():
    doc = Document(SOURCE)

    # Remove the internal checklist page before the cover page.
    for paragraph in list(doc.paragraphs[:25]):
        remove_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        text = paragraph.text

        if text == "FULL CODE: [INSERT GOOGLE DRIVE / COLAB LINK]":
            replace_text(paragraph, "FULL CODE: https://github.com/raichiiiiiii/Machine-Learning")

        elif text.startswith("[TODO: Update with page numbers"):
            replace_text(
                paragraph,
                "Note: Update page numbers using Word's automatic Table of Contents after any final manual formatting changes.",
            )

        elif "[TODO: add one or two sentences summarising" in text:
            replace_text(
                paragraph,
                "The problem domain for this project is established by Sattar et al. (2025), published in MDPI Sustainability, which applied machine learning models to supply chain data. The paper compared XGBoost and recurrent neural networks across supply-chain forecasting and risk-mitigation tasks, reporting strong XGBoost forecasting accuracy and strong RNN performance for fraud and delivery-risk prediction. It also introduced cost-accuracy and ESG-aware evaluation, showing that model selection should consider operational cost and sustainability alongside predictive accuracy. Following the project requirement, this project uses a different dataset (the DataCo Smart Supply Chain dataset) from the one used in the reference paper, while addressing the same domain of supply chain analytics with machine learning.",
            )

        elif "[TODO: confirm the exact list of leakage columns dropped" in text:
            replace_text(
                paragraph,
                "Leakage control was a critical step. Leakage columns are fields that would not be known at prediction time, or that mathematically encode the target. For the fraud task, the order status field that the fraud label is derived from was removed, along with post-event delivery fields and identifier columns. The dropped columns were: Order Status, Delivery Status, Late_delivery_risk, Days for shipping (real), shipping date (DateOrders), Type, Order Id, Order Item Id, Product Card Id, Product Category Id, Customer Id, Department Id, Category Id, Order Customer Id, Customer Email, Customer Password, Customer Fname, Customer Lname, Customer Street, Order Zipcode, Product Description, Product Price, Product Status, Order Item Total, and Order Profit Per Order. Keeping such columns would let a model learn target leakage rather than genuine predictive patterns.",
            )

        elif "[TODO: confirm exact layer sizes" in text:
            replace_text(
                paragraph,
                "Third, the two deep learning models were trained. The MLP is a feed-forward network with Dense(128, ReLU), Batch Normalization, Dropout(0.3), Dense(64, ReLU), Batch Normalization, Dropout(0.3), Dense(32, ReLU), Dropout(0.2), and a Dense(1, sigmoid) output for binary classification. The proposed GRU RNN treats the 11 principal components as a sequence of 11 timesteps with one value each, so the input shape is (samples, timesteps = 11, 1). Its architecture is GRU(64, return_sequences=True), Dropout(0.2), GRU(32), Dropout(0.2), Dense(32, ReLU), Dropout(0.3), and Dense(1, sigmoid). Both deep models used Adam with learning rate 1e-3, binary cross-entropy, class weights, batch size 256, up to 30 epochs, and early stopping on validation PR-AUC.",
            )

        elif "[TODO: confirm this contribution description" in text:
            replace_text(
                paragraph,
                "Alhamdulillah, this project was completed as a combined effort of the group. Br. Azam handled the data analysis component, covering the preprocessing pipeline, ethical data pruning, missing-value and outlier handling, encoding, and the chronological data splits, as well as the dataset description. Br. Harith worked on the classical baseline models, implementing and tuning Logistic Regression and Decision Tree. Br. Alif handled the ensemble baselines, implementing and tuning Random Forest and XGBoost, together with the literature review. Br. Syamil developed the deep learning component, building and training the proposed GRU Recurrent Neural Network and the MLP comparison, and compiling the model evaluation on the PCA feature set. All group members contributed to the abstract, introduction, results discussion, conclusion, and references.",
            )

        elif text == "Sattar et al. (2025). [TODO: insert the full citation of the reference paper, MDPI Sustainability].":
            replace_text(
                paragraph,
                "Sattar, M. U., Dattana, V., Hasan, R., Mahmood, S., Khan, H. W., & Hussain, S. (2025). Enhancing Supply Chain Management: A Comparative Study of Machine Learning Techniques with Cost-Accuracy and ESG-Based Evaluation for Forecasting and Risk Mitigation. Sustainability, 17(13), 5772. https://doi.org/10.3390/su17135772",
            )

        elif text == "[INSERT FIGURE 5: Confusion matrix for Random Forest]":
            insert_image_at_placeholder(paragraph, FIG_DIR / "figure_5_random_forest_confusion.png", width=4.8)
        elif text == "[INSERT FIGURE 6: Confusion matrix for XGBoost]":
            insert_image_at_placeholder(paragraph, FIG_DIR / "figure_6_xgboost_confusion.png", width=4.8)
        elif text == "[INSERT FIGURE 7: MLP training and validation loss curves]":
            insert_image_at_placeholder(paragraph, ROOT / "model" / "MLP" / "figure" / "fraud_pca_95_variance_mlp_loss_curve.png", width=5.4)
        elif text == "[INSERT FIGURE 8: Confusion matrix for MLP]":
            insert_image_at_placeholder(paragraph, ROOT / "model" / "MLP" / "figure" / "fraud_pca_95_variance_mlp_confusion_matrix.png", width=4.8)
        elif text == "[INSERT FIGURE 9: GRU RNN training and validation loss curves]":
            insert_image_at_placeholder(paragraph, ROOT / "model" / "RNN" / "figure" / "fraud_pca_95_variance_rnn_loss_curve.png", width=5.4)
        elif text == "[INSERT FIGURE 10: Confusion matrix for the GRU RNN]":
            insert_image_at_placeholder(paragraph, ROOT / "model" / "RNN" / "figure" / "fraud_pca_95_variance_rnn_confusion_matrix.png", width=4.8)
        elif text == "[INSERT FIGURE 11: PR-AUC and F1-score comparison bar chart for all models]":
            insert_image_at_placeholder(paragraph, FIG_DIR / "figure_11_metric_comparison.png", width=5.8)
        elif text == "[INSERT FIGURE 12: Combined precision-recall and ROC curves for all models]":
            insert_image_at_placeholder(paragraph, FIG_DIR / "figure_12_roc_pr_curves.png", width=5.8)

    doc.save(OUTPUT)


def main():
    generate_figures()
    update_document()
    print(OUTPUT)


if __name__ == "__main__":
    main()
