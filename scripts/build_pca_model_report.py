from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sklearn.metrics import confusion_matrix, precision_recall_curve, roc_curve


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
FIG_DIR = REPORT_DIR / "figures"
REPORT_DIR.mkdir(exist_ok=True)
FIG_DIR.mkdir(parents=True, exist_ok=True)

PCA_DATASET_DEEP = "../PCA/fraud_pca_95_variance.csv"
PCA_DATASET_CLASSICAL = "fraud_pca_95_variance"


def fmt(value, digits=4):
    return f"{float(value):.{digits}f}"


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.6)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths):
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for cell in table.rows[0].cells:
        shade_cell(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(31, 77, 120)


def add_table(doc, dataframe, widths):
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(dataframe.columns):
        set_cell_text(hdr[i], col, bold=True, color="1F4D78")
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(dataframe.columns):
            set_cell_text(cells[i], row[col])
    style_table(table, widths)
    return table


def read_results():
    mlp = pd.read_csv(ROOT / "model" / "MLP" / "csv" / "mlp_results_comparison.csv")
    rnn = pd.read_csv(ROOT / "model" / "RNN" / "csv" / "rnn_results_comparison.csv")
    classical = pd.read_csv(ROOT / "model" / "Classical" / "csv" / "classical_results_comparison.csv")

    rows = []
    for label, df in [("MLP", mlp), ("RNN", rnn)]:
        row = df[df["dataset"] == PCA_DATASET_DEEP].iloc[0].to_dict()
        pred = pd.read_csv(ROOT / "model" / label / "csv" / f"fraud_pca_95_variance_{label.lower()}_test_predictions.csv")
        tn, fp, fn, tp = confusion_matrix(pred["y_true"], pred["y_pred"]).ravel()
        rows.append(
            {
                "model": label,
                "model_family": "Deep learning",
                "features": int(row["n_features"]),
                "threshold": row["best_threshold"],
                "accuracy": row["accuracy"],
                "precision": row["precision"],
                "recall": row["recall"],
                "f1": row["f1"],
                "roc_auc": row["roc_auc"],
                "pr_auc": row["pr_auc"],
                "tn": tn,
                "fp": fp,
                "fn": fn,
                "tp": tp,
            }
        )

    classical_rows = classical[classical["dataset"] == PCA_DATASET_CLASSICAL].copy()
    for _, row in classical_rows.iterrows():
        rows.append(
            {
                "model": row["model"],
                "model_family": "Classical ML",
                "features": int(row["features"]),
                "threshold": row["threshold"],
                "accuracy": row["accuracy"],
                "precision": row["precision"],
                "recall": row["recall"],
                "f1": row["f1"],
                "roc_auc": row["roc_auc"],
                "pr_auc": row["pr_auc"],
                "tn": int(row["tn"]),
                "fp": int(row["fp"]),
                "fn": int(row["fn"]),
                "tp": int(row["tp"]),
            }
        )

    results = pd.DataFrame(rows)
    return results.sort_values("pr_auc", ascending=False).reset_index(drop=True)


def read_probabilities():
    probs = {}
    mlp = pd.read_csv(ROOT / "model" / "MLP" / "csv" / "fraud_pca_95_variance_mlp_test_predictions.csv")
    rnn = pd.read_csv(ROOT / "model" / "RNN" / "csv" / "fraud_pca_95_variance_rnn_test_predictions.csv")
    classical = pd.read_csv(ROOT / "model" / "Classical" / "csv" / "fraud_pca_95_variance_classical_test_predictions.csv")

    y_true = mlp["y_true"].to_numpy()
    probs["MLP"] = mlp["y_prob"].to_numpy()
    probs["RNN"] = rnn["y_prob"].to_numpy()
    probs["Linear Regression"] = classical["linear_regression_probability"].to_numpy()
    probs["Decision Tree"] = classical["decision_tree_probability"].to_numpy()
    probs["Random Forest"] = classical["random_forest_probability"].to_numpy()
    probs["XGBoost"] = classical["xgboost_probability"].to_numpy()
    return y_true, probs


def build_charts(results):
    colors = {
        "Random Forest": "#1F77B4",
        "XGBoost": "#2CA02C",
        "Decision Tree": "#FF7F0E",
        "RNN": "#9467BD",
        "MLP": "#8C564B",
        "Linear Regression": "#7F7F7F",
    }
    sorted_results = results.sort_values("pr_auc", ascending=True)

    fig, axes = plt.subplots(1, 2, figsize=(10.8, 3.8))
    axes[0].barh(sorted_results["model"], sorted_results["pr_auc"], color=[colors[m] for m in sorted_results["model"]])
    axes[0].set_title("PR-AUC on PCA Feature Set")
    axes[0].set_xlabel("PR-AUC")
    axes[0].grid(axis="x", alpha=0.25)

    axes[1].barh(sorted_results["model"], sorted_results["f1"], color=[colors[m] for m in sorted_results["model"]])
    axes[1].set_title("F1-score on PCA Feature Set")
    axes[1].set_xlabel("F1-score")
    axes[1].grid(axis="x", alpha=0.25)

    plt.tight_layout()
    metric_path = FIG_DIR / "pca_metric_comparison.png"
    plt.savefig(metric_path, dpi=180)
    plt.close(fig)

    y_true, probs = read_probabilities()
    fig, axes = plt.subplots(1, 2, figsize=(11, 3.8))
    for model, prob in probs.items():
        precision, recall, _ = precision_recall_curve(y_true, prob)
        fpr, tpr, _ = roc_curve(y_true, prob)
        axes[0].plot(recall, precision, label=model, linewidth=1.7)
        axes[1].plot(fpr, tpr, label=model, linewidth=1.7)
    axes[0].set_title("Precision-Recall Curves")
    axes[0].set_xlabel("Recall")
    axes[0].set_ylabel("Precision")
    axes[0].grid(alpha=0.25)
    axes[0].legend(fontsize=7)
    axes[1].plot([0, 1], [0, 1], "--", color="gray", linewidth=1)
    axes[1].set_title("ROC Curves")
    axes[1].set_xlabel("False Positive Rate")
    axes[1].set_ylabel("True Positive Rate")
    axes[1].grid(alpha=0.25)
    axes[1].legend(fontsize=7)
    plt.tight_layout()
    curve_path = FIG_DIR / "pca_roc_pr_curves_all_models.png"
    plt.savefig(curve_path, dpi=180)
    plt.close(fig)

    return metric_path, curve_path


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    cell.width = Inches(6.35)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 58, 95)
    run.font.size = Pt(10.5)
    set_table_widths(table, [6.35])


def build_report(results, metric_path, curve_path):
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("PCA Model Evaluation and Comparison Report")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.add_run("Dataset: fraud_pca_95_variance.csv | Scope: trained fraud-detection models in this repository").italic = True

    best = results.iloc[0]
    add_callout(
        doc,
        f"Summary: Random Forest is the strongest trained model on the PCA feature set by PR-AUC ({fmt(best['pr_auc'])}) and F1-score ({fmt(best['f1'])}). Overall scores remain low, which suggests PCA compression preserved broad variance but weakened minority-class fraud signal.",
    )

    doc.add_heading("1. Evaluation Scope", level=1)
    p = doc.add_paragraph()
    p.add_run("This report compares every trained model available for the PCA dataset. ").bold = True
    p.add_run(
        "The PCA dataset is treated as the feature-selected and PCA-analyzed representation of the fraud pipeline. "
        "It contains 11 principal-component features plus the binary fraud target. The same repository artifacts were used for the report; no model was retrained while compiling this document."
    )

    scope = pd.DataFrame(
        [
            ["Dataset", "PCA/fraud_pca_95_variance.csv"],
            ["Feature representation", "Wrapper-selected features transformed with PCA at 95% variance retention"],
            ["Target", "fraud (1 = suspected fraud, 0 = non-fraud)"],
            ["Evaluated model families", "MLP, RNN, Linear Regression, Decision Tree, Random Forest, XGBoost"],
            ["Primary ranking metric", "PR-AUC, because the fraud class is rare"],
        ],
        columns=["Item", "Value"],
    )
    add_table(doc, scope, [1.9, 4.45])

    doc.add_heading("2. Model Evaluation Results", level=1)
    display_results = results.copy()
    display_results = display_results[
        ["model", "model_family", "features", "threshold", "accuracy", "precision", "recall", "f1", "roc_auc", "pr_auc"]
    ]
    display_results.columns = ["Model", "Family", "Features", "Threshold", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC", "PR-AUC"]
    for col in ["Threshold", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC", "PR-AUC"]:
        display_results[col] = display_results[col].map(lambda x: fmt(x))
    add_table(doc, display_results, [1.25, 1.05, 0.55, 0.7, 0.7, 0.75, 0.65, 0.55, 0.7, 0.7])

    doc.add_paragraph(
        "Interpretation: accuracy is high for several tree models because the non-fraud class dominates. PR-AUC, recall, and F1 give a more useful view of fraud detection quality."
    )

    doc.add_heading("3. Confusion Matrix Summary", level=1)
    cm = results[["model", "tn", "fp", "fn", "tp"]].copy()
    cm.columns = ["Model", "TN", "FP", "FN", "TP"]
    add_table(doc, cm, [1.7, 0.85, 0.85, 0.85, 0.85])

    doc.add_heading("4. Visual Comparison", level=1)
    doc.add_paragraph("The following chart compares the two most relevant aggregate metrics for rare-event fraud classification.")
    doc.add_picture(str(metric_path), width=Inches(5.75))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 1. PR-AUC and F1-score comparison for all trained models on the PCA feature set.")

    doc.add_paragraph("The combined PR and ROC curves show how the model probability scores behave across thresholds.")
    doc.add_picture(str(curve_path), width=Inches(5.75))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Figure 2. Precision-recall and ROC curves for PCA-trained models.")

    doc.add_heading("5. Findings", level=1)
    findings = [
        "Random Forest ranks first by both PR-AUC and F1-score on the PCA representation.",
        "XGBoost is the second strongest classical model by PR-AUC, but its recall and F1 remain modest.",
        "Decision Tree reaches the highest recall among the top three classical models, but it also creates many false positives.",
        "MLP and RNN perform close to the linear baseline on this PCA feature set, indicating that the PCA representation does not expose strong nonlinear fraud patterns to the neural models.",
        "Because fraud is rare, a model can show high accuracy while still missing most fraud cases. PR-AUC and F1 should be prioritized over accuracy in the final comparison.",
    ]
    for item in findings:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        "For the PCA-analyzed fraud dataset, Random Forest is the recommended baseline among the currently trained models. "
        "However, all PCA-based scores are weak compared with the prepared numeric feature-space results already present in the repository, so the PCA representation should be treated as a compact experimental view rather than the best production feature set."
    )

    doc.add_section(WD_SECTION_START.CONTINUOUS)
    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Generated from repository artifacts in model/*/csv").font.size = Pt(8)

    out = REPORT_DIR / "pca_model_evaluation_report.docx"
    doc.save(out)
    return out


def main():
    results = read_results()
    results.to_csv(REPORT_DIR / "pca_model_evaluation_summary.csv", index=False)
    metric_path, curve_path = build_charts(results)
    out = build_report(results, metric_path, curve_path)
    print(out)


if __name__ == "__main__":
    main()
